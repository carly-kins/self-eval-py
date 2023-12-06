#!/usr/bin/env python3

import shutil
import selfeval_vars
import selfeval_process

from pyutils import now_utc
from docx import Document
from docx.shared import Pt

def print_result_table(rows):
    terminal_width = shutil.get_terminal_size().columns
    for terminal_preset_width in sorted(selfeval_vars.WIDTH_PRESETS.keys()):
        if terminal_width <= terminal_preset_width:
            preset = selfeval_vars.WIDTH_PRESETS[terminal_preset_width]
            break
    LG_WIDTH, MD_WIDTH, SM_WIDTH = preset
    XS_WIDTH = 4
    def cell(cell_text, max_width):
        return_cell = cell_text[:max_width-3]
        if len(cell_text) <= max_width:
            return cell_text
        else:
            if cell_text != date_added:
                return return_cell + '...'
            else:
                return return_cell

    fill_dashes_lg = '─' * (LG_WIDTH + 2)
    fill_dashes_md = '─' * (MD_WIDTH + 2)
    fill_dashes_sm = '─' * (SM_WIDTH + 2)
    fill_dashes_xs = '─' * (XS_WIDTH + 2)

    print(f'┌──────┬' + fill_dashes_xs + '┬' + fill_dashes_lg + '┬' + fill_dashes_lg + '┬' + fill_dashes_md + '┬' + fill_dashes_sm + '┐')
    print(f'│  ID  │ {"Type":^{XS_WIDTH}} │ {"Entry":^{LG_WIDTH}} │ {"Tags":^{LG_WIDTH}} │ {"Notes":^{MD_WIDTH}} │ {"Date":^{SM_WIDTH}} │')
    print(f'├──────┼' + fill_dashes_xs + '┼' + fill_dashes_lg + '┼' + fill_dashes_lg + '┼' + fill_dashes_md + '┼' + fill_dashes_sm + '┤')

    for idx, row in enumerate(rows):
        id, entry, types, notes, tags, date_added = row
        if types == "Achievement": types = "A" 
        else : types = "I"
        types, entry, tags, notes, date_added = (cell(types, SM_WIDTH), cell(entry, LG_WIDTH), cell(tags, LG_WIDTH), cell(notes, MD_WIDTH), cell(date_added, SM_WIDTH))
        print(f'│ {id:^{XS_WIDTH}} │ {types:^{XS_WIDTH}} │ {entry:{LG_WIDTH}} │ {tags:{LG_WIDTH}} │ {notes:{MD_WIDTH}} │ {date_added:{SM_WIDTH}} │')
        if idx == len(rows) - 1:
            print(f'└──────┴' + fill_dashes_xs + '┴' + fill_dashes_lg + '┴' + fill_dashes_lg + '┴' + fill_dashes_md + '┴' + fill_dashes_sm + '┘')
        else:
            print(f'├──────┼' + fill_dashes_xs + '┼' + fill_dashes_lg + '┼' +  fill_dashes_lg + '┼' + fill_dashes_md + '┼' + fill_dashes_sm + '┤')

def csv_output(rows):
    for row in rows:
        (id, entry, types, notes, tags, date_added) = row
        print(f'"{id}","{entry}","{types}","{notes}","{tags}","{date_added}"')

def doc_output(rows):

    print(f"Great work! Generating your evaluation...")
    # Generate the document
    document = Document()
    # Set styles
    style = document.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    #get database entries for each of the tables
    con, cur = selfeval_process.create_or_open_db()
    tags_rows = cur.execute(f' SELECT * FROM {selfeval_vars.TAGS_TABLE};').fetchall()
    types_rows = cur.execute(f' SELECT * FROM {selfeval_vars.TYPES_TABLE};').fetchall()
    summary_rows = cur.execute(f' SELECT * FROM {selfeval_vars.SUMMARY_TABLE};').fetchall()
    questions_rows = cur.execute(f' SELECT * FROM {selfeval_vars.QUESTIONS_TABLE};').fetchall()
    con.close()

    def header(text):
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run(text).italic = True

    def bullet_row(rows, check, types_check, include_notes): 
        for row in rows:
            (_, entry, types, notes, tags, _) = row

            if types_check: # look at that cognitive complexity :snap_point: 
                if check in types:
                    document.add_paragraph(entry, style='List Bullet')
                    if include_notes:
                        if notes: document.add_paragraph(notes, style='List Bullet 2')
            else:
                if check in tags:
                    document.add_paragraph(entry, style='List Bullet')
                    if notes: document.add_paragraph(notes, style='List Bullet 2')

    # TITLE 
    document.add_heading(f'{selfeval_vars.NAME} - Self Evaluation', 0)
    document.add_paragraph(f'Generated on {now_utc()}', style='Subtitle').paragraph_format.space_after = Pt(10)

    ## MONTHLY EVAL
    document.add_heading('Monthly Evaluation', level=1)

    for questions_row in questions_rows:
        ( q_id, q_title, q_desc, type_id ) = questions_row

        for types_row in types_rows:
            ( id, type ) = types_row
            if (id == type_id): 
                document.add_heading(q_title, level=2)
                header(q_desc)
                if q_id == 2: include_notes = False
                else: include_notes = True
                bullet_row(rows, type, True, include_notes)

    document.add_page_break()

    ## EOY EVAL
    document.add_heading('End of Year Evaluation', level=1)

    for summary_row in summary_rows:
        ( _, summary ) = summary_row

        count = 0
        for tags_row in tags_rows:
            ( _, tag, desc, tag_summary ) = tags_row
            if (summary == tag_summary) & (count < 1): 
                document.add_heading(tag_summary, level=2)
                count +=1
            if (summary == tag_summary):
                document.add_heading(tag, level=3)
                header(desc)
                bullet_row(rows, tag, False, True)

    print(f"Saved {selfeval_vars.NAME}-self-eval-{now_utc()}.docx to your desktop!")
    document.save(f'{selfeval_vars.HOME}/Desktop/{selfeval_vars.NAME}-self-eval-{now_utc()}.docx')